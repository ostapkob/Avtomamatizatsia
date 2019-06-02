import docx
import os

#===========READER==============
doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))
for i in range(len(doc.paragraphs)):
    pass
    # print(doc.paragraphs[i].text)

doc.paragraphs[0].style = 'Title'
print(doc.paragraphs[1].runs[3].text)
doc.paragraphs[1].runs[2].underline = True

doc.save('DocTest.docx')
# os.system('DocTest.docx')

#===========WRITER==============
doc = docx.Document()
doc.add_paragraph('Hello world')
Obj1 = doc.add_paragraph('I am Ostap')
Obj2 = doc.add_paragraph('I love Yulya')
Obj1.add_run('-yyyyyyyyyyyyy')
doc.paragraphs[2].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

text = 'abcd'
content = 'dfgdfg','dfg','uyiyt','sdrdwe','cxzcvsdcv','sdfsdf','dfgdfg','dfgdfg','dfgdfghyj'

for i, t in enumerate(text):
    doc.add_heading(t, i)
    doc.add_paragraph('Hello world')

doc.save('DocNew.docx')
os.system('DocNew.docx')




